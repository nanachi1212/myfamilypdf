from io import BytesIO
import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image
from pypdf import PdfReader, PdfWriter, Transformation
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from familypdf_office_export.docx_writer import write_docx
from familypdf_office_export.extract import extract_document


def _merge_raster_image(
    page,
    size: tuple[int, int],
    color: tuple[int, int, int],
    x: float,
    y: float,
) -> None:
    image = Image.new("RGB", size, color)
    image_pdf = BytesIO()
    image.save(image_pdf, format="PDF", resolution=72.0)
    image_pdf.seek(0)
    image_page = PdfReader(image_pdf).pages[0]
    page.merge_transformed_page(
        image_page,
        Transformation().translate(x, y),
    )


def _write_two_column_pdf(path: Path) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=600, height=400)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): font_reference}
            )
        }
    )
    content = DecodedStreamObject()
    content.set_data(
        b"""
BT /F1 18 Tf 220 375 Td (Full width heading) Tj ET
BT /F1 12 Tf 40 340 Td (Left top) Tj ET
BT /F1 12 Tf 40 300 Td (Left bottom) Tj ET
BT /F1 12 Tf 340 340 Td (Right top) Tj ET
BT /F1 12 Tf 340 300 Td (Right bottom) Tj ET
BT /F1 16 Tf 220 250 Td (Middle heading) Tj ET
BT /F1 12 Tf 40 210 Td (Left second top) Tj ET
BT /F1 12 Tf 40 170 Td (Left second bottom) Tj ET
BT /F1 12 Tf 340 210 Td (Right second top) Tj ET
BT /F1 12 Tf 340 170 Td (Right second bottom) Tj ET
"""
    )
    page[NameObject("/Contents")] = writer._add_object(content)
    _merge_raster_image(page, (100, 20), (30, 90, 220), 80, 320)
    _merge_raster_image(page, (120, 40), (220, 40, 60), 240, 80)
    with path.open("wb") as stream:
        writer.write(stream)


def _write_single_column_image_pdf(path: Path) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=600, height=400)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): font_reference}
            )
        }
    )
    content = DecodedStreamObject()
    content.set_data(
        b"""
BT /F1 12 Tf 40 340 Td (Before image) Tj ET
BT /F1 12 Tf 40 240 Td (After image) Tj ET
"""
    )
    page[NameObject("/Contents")] = writer._add_object(content)
    _merge_raster_image(page, (80, 40), (30, 180, 90), 40, 280)
    with path.open("wb") as stream:
        writer.write(stream)


class MultiColumnExportTest(unittest.TestCase):
    def test_preserves_single_column_raster_image_order(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "single-column-image.pdf"
            target = Path(directory) / "single-column-image.docx"
            _write_single_column_image_pdf(source)

            extracted = extract_document(source)
            report = write_docx(extracted, target)
            document = Document(target)

        self.assertEqual(extracted.pages[0].column_count, 1)
        self.assertEqual(
            [
                item.text if hasattr(item, "text") else "<image>"
                for item in extracted.pages[0].layout_items
            ],
            ["Before image", "<image>", "After image"],
        )
        self.assertEqual(len(document.tables), 0)
        self.assertEqual(len(document.inline_shapes), 1)
        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs],
            ["Before image", "", "After image"],
        )
        self.assertEqual(report.paragraphs_exported, 2)
        self.assertEqual(report.images_exported, 1)

    def test_preserves_two_pdf_columns_as_editable_docx_columns(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "two-columns.pdf"
            target = Path(directory) / "two-columns.docx"
            _write_two_column_pdf(source)

            extracted = extract_document(source)
            report = write_docx(extracted, target)
            document = Document(target)

        self.assertEqual(extracted.pages[0].column_count, 2)
        self.assertEqual(len(extracted.pages[0].images), 2)
        left_image, spanning_image = extracted.pages[0].images
        self.assertEqual((left_image.column, left_image.column_span), (0, 1))
        self.assertEqual(
            (spanning_image.column, spanning_image.column_span),
            (0, 2),
        )
        self.assertTrue(left_image.data.startswith(b"\x89PNG\r\n\x1a\n"))
        self.assertEqual(
            (left_image.width_points, left_image.height_points),
            (100.0, 20.0),
        )
        self.assertEqual(
            [
                item.text if hasattr(item, "text") else "<image>"
                for item in extracted.pages[0].layout_items
            ],
            [
                "Full width heading",
                "Left top",
                "<image>",
                "Left bottom",
                "Right top",
                "Right bottom",
                "Middle heading",
                "Left second top",
                "Left second bottom",
                "Right second top",
                "Right second bottom",
                "<image>",
            ],
        )
        heading = extracted.pages[0].blocks[0]
        self.assertEqual(heading.text, "Full width heading")
        self.assertEqual(heading.column_span, 2)
        self.assertEqual(
            [block.text for block in extracted.pages[0].blocks],
            [
                "Full width heading",
                "Left top",
                "Left bottom",
                "Right top",
                "Right bottom",
                "Middle heading",
                "Left second top",
                "Left second bottom",
                "Right second top",
                "Right second bottom",
            ],
        )
        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            ["Full width heading", "Middle heading"],
        )
        body_order = [
            child.tag.rsplit("}", 1)[-1]
            for child in document.element.body.iterchildren()
            if child.tag.rsplit("}", 1)[-1] in {"p", "tbl"}
        ]
        self.assertEqual(body_order, ["p", "tbl", "p", "tbl", "p"])
        self.assertEqual(len(document.inline_shapes), 2)
        self.assertEqual(len(document.tables), 2)
        first_cells = document.tables[0].rows[0].cells
        self.assertEqual(len(first_cells), 2)
        self.assertEqual(
            [paragraph.text for paragraph in first_cells[0].paragraphs],
            ["Left top", "", "Left bottom"],
        )
        self.assertEqual(
            first_cells[0].text.splitlines(),
            ["Left top", "", "Left bottom"],
        )
        self.assertEqual(
            first_cells[1].text.splitlines(),
            ["Right top", "Right bottom"],
        )
        second_cells = document.tables[1].rows[0].cells
        self.assertEqual(
            second_cells[0].text.splitlines(),
            ["Left second top", "Left second bottom"],
        )
        self.assertEqual(
            second_cells[1].text.splitlines(),
            ["Right second top", "Right second bottom"],
        )
        self.assertEqual(report.paragraphs_exported, 10)
        self.assertEqual(report.images_exported, 2)


if __name__ == "__main__":
    unittest.main()
