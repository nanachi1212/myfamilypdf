import contextlib
import io
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from familypdf_office_export.cli import (
    EXIT_MISSING_TEXT_LAYER,
    parse_page_range,
    run,
)


def _write_two_page_pdf(path: Path) -> None:
    writer = PdfWriter()
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)

    for label in (b"First page", b"Second page"):
        page = writer.add_blank_page(width=300, height=300)
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): font_reference}
                )
            }
        )
        content = DecodedStreamObject()
        content.set_data(
            b"BT /F1 12 Tf 40 260 Td (" + label + b") Tj ET"
        )
        page[NameObject("/Contents")] = writer._add_object(content)

    with path.open("wb") as stream:
        writer.write(stream)


class CliTest(unittest.TestCase):
    def test_parse_page_range_keeps_order_and_removes_duplicates(self) -> None:
        self.assertEqual(parse_page_range("3,1-2,2,5"), [3, 1, 2, 5])
        self.assertIsNone(parse_page_range(""))
        with self.assertRaisesRegex(ValueError, "Invalid page range"):
            parse_page_range("3-1")

    def test_exports_selected_page_to_docx_and_prints_json_report(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "source.pdf"
            target = Path(directory) / "selected.docx"
            _write_two_page_pdf(source)
            stdout = io.StringIO()

            with contextlib.redirect_stdout(stdout):
                exit_code = run(
                    [
                        "--input",
                        str(source),
                        "--output",
                        str(target),
                        "--format",
                        "docx",
                        "--pages",
                        "2",
                    ]
                )

            report = json.loads(stdout.getvalue())
            document = Document(target)

        self.assertEqual(exit_code, 0)
        self.assertEqual(report["status"], "ok")
        self.assertEqual(report["pages_exported"], 1)
        self.assertEqual(report["images_exported"], 0)
        self.assertIn("Second page", "\n".join(p.text for p in document.paragraphs))
        self.assertNotIn("First page", "\n".join(p.text for p in document.paragraphs))

    def test_rejects_document_without_searchable_text_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "blank.pdf"
            target = Path(directory) / "blank.xlsx"
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=300)
            with source.open("wb") as stream:
                writer.write(stream)
            stdout = io.StringIO()

            with contextlib.redirect_stdout(stdout):
                exit_code = run(
                    [
                        "--input",
                        str(source),
                        "--output",
                        str(target),
                        "--format",
                        "xlsx",
                    ]
                )
            report = json.loads(stdout.getvalue())

        self.assertEqual(exit_code, EXIT_MISSING_TEXT_LAYER)
        self.assertEqual(report["status"], "needs_ocr")
        self.assertFalse(target.exists())


if __name__ == "__main__":
    unittest.main()
