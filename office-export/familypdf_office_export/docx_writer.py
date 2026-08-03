from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Emu, Pt

from .model import (
    ExtractedDocument,
    ExtractedImage,
    LayoutItem,
    TextBlock,
)


def _write_block(paragraph, block) -> None:
    for source_run in block.runs:
        run = paragraph.add_run(source_run.text)
        run.bold = source_run.bold
        run.italic = source_run.italic
        if source_run.font_size is not None:
            run.font.size = Pt(source_run.font_size)
        if source_run.font_name:
            run.font.name = source_run.font_name


def _write_image(
    paragraph,
    image: ExtractedImage,
    max_width_emu: int,
) -> None:
    width_emu = min(int(Pt(image.width_points)), max_width_emu)
    paragraph.add_run().add_picture(
        BytesIO(image.data),
        width=Emu(width_emu),
    )


def _write_column_table(
    document,
    columns: list[list[LayoutItem]],
) -> tuple[int, int]:
    if not any(columns):
        return 0, 0

    table = document.add_table(rows=1, cols=len(columns))
    paragraphs_exported = 0
    images_exported = 0
    section = document.sections[-1]
    usable_width_emu = int(
        section.page_width
        - section.left_margin
        - section.right_margin
    )
    column_width_emu = int(usable_width_emu / len(columns) * 0.9)
    for column_index, cell in enumerate(table.rows[0].cells):
        for item_index, item in enumerate(columns[column_index]):
            paragraph = (
                cell.paragraphs[0]
                if item_index == 0
                else cell.add_paragraph()
            )
            if isinstance(item, TextBlock):
                _write_block(paragraph, item)
                paragraphs_exported += 1
            else:
                _write_image(paragraph, item, column_width_emu)
                images_exported += 1
    return paragraphs_exported, images_exported


@dataclass(slots=True, frozen=True)
class DocxExportReport:
    pages_exported: int
    paragraphs_exported: int
    images_exported: int


def write_docx(
    extracted: ExtractedDocument,
    output_path: str | Path,
) -> DocxExportReport:
    """Write editable paragraphs and basic run styles to a DOCX file."""
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    paragraphs_exported = 0
    images_exported = 0
    for page_index, page in enumerate(extracted.pages):
        layout_items = page.layout_items or list(page.blocks)
        if page.column_count > 1:
            pending_columns: list[list[LayoutItem]] = [
                [] for _ in range(page.column_count)
            ]
            for item in layout_items:
                if item.column_span > 1:
                    table_paragraphs, table_images = _write_column_table(
                        document,
                        pending_columns,
                    )
                    paragraphs_exported += table_paragraphs
                    images_exported += table_images
                    pending_columns = [
                        [] for _ in range(page.column_count)
                    ]
                    paragraph = document.add_paragraph()
                    if isinstance(item, TextBlock):
                        _write_block(paragraph, item)
                        paragraphs_exported += 1
                    else:
                        section = document.sections[-1]
                        usable_width_emu = int(
                            section.page_width
                            - section.left_margin
                            - section.right_margin
                        )
                        _write_image(paragraph, item, usable_width_emu)
                        images_exported += 1
                else:
                    pending_columns[item.column].append(item)
            table_paragraphs, table_images = _write_column_table(
                document,
                pending_columns,
            )
            paragraphs_exported += table_paragraphs
            images_exported += table_images
        else:
            section = document.sections[-1]
            usable_width_emu = int(
                section.page_width
                - section.left_margin
                - section.right_margin
            )
            for item in layout_items:
                paragraph = document.add_paragraph()
                if isinstance(item, TextBlock):
                    _write_block(paragraph, item)
                    paragraphs_exported += 1
                else:
                    _write_image(paragraph, item, usable_width_emu)
                    images_exported += 1

        if page_index + 1 < len(extracted.pages):
            document.add_page_break()

    document.save(target)
    return DocxExportReport(
        pages_exported=len(extracted.pages),
        paragraphs_exported=paragraphs_exported,
        images_exported=images_exported,
    )
